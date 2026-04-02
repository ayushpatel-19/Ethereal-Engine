"""
Ethereal Engine — Ingestion Pipeline
Phase 01: Data + Ingestion

Handles:
  • PDF (text + OCR fallback)
  • DOCX
  • Plain text
  • Web URLs (static + JS-rendered via Playwright)
  • REST APIs
  • Preprocessing: cleaning, deduplication, structure detection, PII masking
"""
from __future__ import annotations

import re
import uuid
import hashlib
import asyncio
from pathlib import Path
from typing import AsyncGenerator

import httpx
import pdfplumber
import pytesseract
from PIL import Image
from pdf2image import convert_from_path
from pytesseract import TesseractNotFoundError
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import trafilatura
from loguru import logger

from core.config import get_settings
from core.models import RawDocument, DocumentMetadata, SourceType

settings = get_settings()
if settings.tesseract_cmd_resolved:
    pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd_resolved


# ══════════════════════════════════════════════════════════════════════════════
# PII Masking patterns
# ══════════════════════════════════════════════════════════════════════════════
PII_PATTERNS = [
    (re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'), "[EMAIL]"),
    (re.compile(r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b'), "[PHONE]"),
    (re.compile(r'\b\d{3}-\d{2}-\d{4}\b'), "[SSN]"),
    (re.compile(r'\b(?:\d[ -]*?){13,16}\b'), "[CARD_NUMBER]"),
    (re.compile(r'\b[A-Z]{2}\d{6}[A-Z]?\b'), "[PASSPORT]"),
]


def mask_pii(text: str) -> str:
    """Redact personally identifiable information from text."""
    for pattern, replacement in PII_PATTERNS:
        text = pattern.sub(replacement, text)
    return text


def clean_text(text: str) -> str:
    """Normalize whitespace, remove control characters, fix encoding artifacts."""
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    text = re.sub(r'-\n([a-z])', r'\1', text)  # Fix hyphenated line breaks
    return text.strip()


def deduplicate_paragraphs(text: str) -> str:
    """Remove duplicate paragraphs (common in scanned docs)."""
    paragraphs = text.split('\n\n')
    seen = set()
    unique = []
    for p in paragraphs:
        fingerprint = hashlib.md5(p.strip().lower().encode()).hexdigest()
        if fingerprint not in seen:
            seen.add(fingerprint)
            unique.append(p)
    return '\n\n'.join(unique)


def detect_structure(text: str) -> dict:
    """Heuristically detect document structure: headings, tables, lists."""
    lines = text.split('\n')
    has_headings = any(re.match(r'^#{1,6}\s|^[A-Z][A-Z\s]{4,}$', l.strip()) for l in lines)
    has_tables   = '|' in text and text.count('|') > 5
    has_lists    = any(re.match(r'^\s*[-•*]\s|^\s*\d+\.\s', l) for l in lines)
    return {"has_headings": has_headings, "has_tables": has_tables, "has_lists": has_lists}


# ══════════════════════════════════════════════════════════════════════════════
# PDF Ingestion
# ══════════════════════════════════════════════════════════════════════════════

async def ingest_pdf(file_path: str | Path) -> RawDocument:
    """Extract text from PDF. Falls back to OCR if text layer is absent."""
    path = Path(file_path)
    logger.info(f"Ingesting PDF: {path.name}")

    text_parts = []
    page_count = 0
    ocr_required_pages: list[int] = []
    ocr_skipped_pages: list[int] = []
    ocr_failed_pages: list[int] = []
    tesseract_available = bool(settings.tesseract_cmd_resolved)

    with pdfplumber.open(path) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            extracted = page.extract_text()
            if extracted and len(extracted.strip()) > 20:
                text_parts.append(extracted)
            else:
                ocr_required_pages.append(page.page_number)
                logger.debug(f"OCR fallback for page {page.page_number}")
                if not tesseract_available:
                    ocr_skipped_pages.append(page.page_number)
                    continue

                try:
                    images = convert_from_path(
                        str(path),
                        first_page=page.page_number,
                        last_page=page.page_number,
                        dpi=300,
                    )
                    if images:
                        ocr_text = pytesseract.image_to_string(images[0])
                        if ocr_text and ocr_text.strip():
                            text_parts.append(ocr_text)
                        else:
                            ocr_failed_pages.append(page.page_number)
                    else:
                        ocr_failed_pages.append(page.page_number)
                except TesseractNotFoundError:
                    tesseract_available = False
                    ocr_skipped_pages.append(page.page_number)
                except Exception as exc:
                    logger.warning(f"OCR failed for page {page.page_number}: {exc}")
                    ocr_failed_pages.append(page.page_number)

    raw_text = '\n\n'.join(text_parts)
    raw_text = clean_text(deduplicate_paragraphs(raw_text))
    raw_text = mask_pii(raw_text)

    if ocr_skipped_pages:
        if raw_text:
            logger.warning(
                "Skipped OCR for page(s) {} because Tesseract is unavailable. "
                "Continuing with embedded PDF text from the remaining pages.".format(
                    ", ".join(map(str, ocr_skipped_pages))
                )
            )
        else:
            raise RuntimeError(
                "This PDF appears to be scanned and requires OCR, but Tesseract is not installed. "
                "Install Tesseract OCR and add it to PATH, or set TESSERACT_CMD in backend/.env."
            )

    if ocr_failed_pages:
        logger.warning(
            "OCR did not extract text for page(s) {}. Continuing with the text that was available.".format(
                ", ".join(map(str, ocr_failed_pages))
            )
        )

    return RawDocument(
        id=str(uuid.uuid4()),
        content=raw_text,
        metadata=DocumentMetadata(
            source=str(path),
            source_type=SourceType.PDF,
            title=path.stem,
            page_count=page_count,
        )
    )


# ══════════════════════════════════════════════════════════════════════════════
# DOCX Ingestion
# ══════════════════════════════════════════════════════════════════════════════

async def ingest_docx(file_path: str | Path) -> RawDocument:
    """Extract text from Word (.docx) files, including tables."""
    path = Path(file_path)
    logger.info(f"Ingesting DOCX: {path.name}")

    doc = DocxDocument(str(path))
    parts = []

    for element in doc.element.body:
        tag = element.tag.split('}')[-1]
        if tag == 'p':
            text = ''.join(node.text or '' for node in element.iter() if node.text)
            if text.strip():
                parts.append(text.strip())
        elif tag == 'tbl':
            # Extract table as markdown
            rows = []
            for row in element.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr'):
                cells = [
                    ''.join(node.text or '' for node in cell.iter() if node.text)
                    for cell in row.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')
                ]
                rows.append(' | '.join(cells))
            if rows:
                parts.append('\n'.join(rows))

    raw_text = clean_text(mask_pii('\n\n'.join(parts)))

    # Extract author from doc properties
    author = None
    if doc.core_properties.author:
        author = doc.core_properties.author

    return RawDocument(
        id=str(uuid.uuid4()),
        content=raw_text,
        metadata=DocumentMetadata(
            source=str(path),
            source_type=SourceType.DOCX,
            title=doc.core_properties.title or path.stem,
            author=author,
        )
    )


# ══════════════════════════════════════════════════════════════════════════════
# Plain Text Ingestion
# ══════════════════════════════════════════════════════════════════════════════

async def ingest_text(file_path: str | Path) -> RawDocument:
    path = Path(file_path)
    logger.info(f"Ingesting text: {path.name}")
    content = path.read_text(encoding='utf-8', errors='replace')
    content = clean_text(mask_pii(content))
    return RawDocument(
        id=str(uuid.uuid4()),
        content=content,
        metadata=DocumentMetadata(
            source=str(path),
            source_type=SourceType.TXT,
            title=path.stem,
        )
    )


# ══════════════════════════════════════════════════════════════════════════════
# Web URL Ingestion
# ══════════════════════════════════════════════════════════════════════════════

async def ingest_url(url: str, max_depth: int = 1, max_pages: int = 10,
                     _visited: set | None = None, _depth: int = 0
                     ) -> AsyncGenerator[RawDocument, None]:
    """Crawl and extract content from a URL and its linked pages."""
    if _visited is None:
        _visited = set()

    if url in _visited or len(_visited) >= max_pages or _depth > max_depth:
        return

    _visited.add(url)
    logger.info(f"Crawling [{_depth}]: {url}")

    try:
        async with httpx.AsyncClient(timeout=30, follow_redirects=True) as client:
            response = await client.get(url, headers={"User-Agent": "EtherealEngine/2.4 (+bot)"})
            response.raise_for_status()
            html = response.text

        # Use trafilatura for clean article extraction
        extracted = trafilatura.extract(
            html,
            include_comments=False,
            include_tables=True,
            no_fallback=False,
        )

        if not extracted:
            # Fallback: BeautifulSoup
            soup = BeautifulSoup(html, 'html.parser')
            for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
                tag.decompose()
            extracted = soup.get_text(separator='\n', strip=True)

        if extracted and len(extracted.strip()) > 100:
            # Get title
            soup = BeautifulSoup(html, 'html.parser')
            title = soup.find('title')
            title_text = title.get_text(strip=True) if title else url

            cleaned = clean_text(mask_pii(extracted))
            yield RawDocument(
                id=str(uuid.uuid4()),
                content=cleaned,
                metadata=DocumentMetadata(
                    source=url,
                    source_type=SourceType.URL,
                    title=title_text,
                )
            )

        # Crawl linked pages if depth allows
        if _depth < max_depth:
            soup = BeautifulSoup(html, 'html.parser')
            from urllib.parse import urljoin, urlparse
            base_domain = urlparse(url).netloc
            links = set()
            for a in soup.find_all('a', href=True):
                href = urljoin(url, a['href'])
                if urlparse(href).netloc == base_domain and href not in _visited:
                    links.add(href)

            for link in list(links)[:5]:  # Max 5 links per page
                async for doc in ingest_url(link, max_depth, max_pages, _visited, _depth + 1):
                    yield doc
                await asyncio.sleep(0.5)  # Polite crawling delay

    except Exception as e:
        logger.warning(f"Failed to crawl {url}: {e}")


# ══════════════════════════════════════════════════════════════════════════════
# REST API Ingestion
# ══════════════════════════════════════════════════════════════════════════════

async def ingest_api(endpoint: str, method: str = "GET",
                     headers: dict | None = None,
                     body: dict | None = None,
                     json_path: str | None = None) -> RawDocument:
    """Fetch data from a REST API and extract text content."""
    logger.info(f"Ingesting API: {method} {endpoint}")
    headers = headers or {}

    async with httpx.AsyncClient(timeout=30) as client:
        if method.upper() == "GET":
            response = await client.get(endpoint, headers=headers)
        elif method.upper() == "POST":
            response = await client.post(endpoint, headers=headers, json=body)
        else:
            raise ValueError(f"Unsupported method: {method}")

        response.raise_for_status()
        data = response.json()

    # Extract text using simple JSONPath-like traversal
    text_content = _extract_json_text(data, json_path)
    cleaned = clean_text(mask_pii(text_content))

    return RawDocument(
        id=str(uuid.uuid4()),
        content=cleaned,
        metadata=DocumentMetadata(
            source=endpoint,
            source_type=SourceType.API,
            title=f"API: {endpoint}",
        )
    )


def _extract_json_text(data: any, json_path: str | None = None) -> str:
    """Recursively extract all string values from JSON."""
    if json_path:
        # Simple dot-notation path support: "data.items"
        keys = json_path.strip("$.").split(".")
        for key in keys:
            if isinstance(data, dict):
                data = data.get(key, {})
            elif isinstance(data, list):
                try:
                    data = data[int(key)]
                except (ValueError, IndexError):
                    break

    def _flatten(obj, parts=None):
        if parts is None:
            parts = []
        if isinstance(obj, str):
            parts.append(obj)
        elif isinstance(obj, dict):
            for v in obj.values():
                _flatten(v, parts)
        elif isinstance(obj, list):
            for item in obj:
                _flatten(item, parts)
        return parts

    return '\n\n'.join(_flatten(data))


# ══════════════════════════════════════════════════════════════════════════════
# Universal Dispatcher
# ══════════════════════════════════════════════════════════════════════════════

async def ingest_file(file_path: str | Path) -> RawDocument:
    """Route a file to the correct ingestion handler based on extension."""
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == '.pdf':
        return await ingest_pdf(path)
    elif ext == '.docx':
        return await ingest_docx(path)
    elif ext in ('.txt', '.md', '.rst', '.csv'):
        return await ingest_text(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
